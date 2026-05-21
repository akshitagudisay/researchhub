import io
import json
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db
from .. import models
from ..auth import get_current_user

router = APIRouter(prefix="/export", tags=["export"])

SECTIONS_ORDER = ["abstract", "introduction", "methodology", "results", "conclusion"]
SECTION_LABELS = {
    "abstract": "Abstract",
    "introduction": "Introduction",
    "methodology": "Methodology",
    "results": "Results",
    "conclusion": "Conclusion",
}


def _assert_access(project_id: int, user: models.User, db: Session) -> models.Project:
    project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    is_owner = project.owner_id == user.id
    is_collab = db.query(models.Collaborator).filter(
        models.Collaborator.project_id == project_id,
        models.Collaborator.user_id == user.id,
    ).first() is not None
    if not is_owner and not is_collab:
        raise HTTPException(status_code=403, detail="Not authorized")
    return project


def _load_data(project_id: int, db: Session):
    manuscript = db.query(models.Manuscript).filter(
        models.Manuscript.project_id == project_id
    ).first()
    content: dict = {}
    if manuscript:
        try:
            content = json.loads(manuscript.content)
        except Exception:
            content = {}

    citations = (
        db.query(models.Citation)
        .filter(models.Citation.project_id == project_id)
        .order_by(models.Citation.created_at.asc())
        .all()
    )
    owner = db.query(models.User).filter(
        models.User.id == db.query(models.Project).filter(
            models.Project.id == project_id
        ).first().owner_id
    ).first()
    collaborators = db.query(models.Collaborator).filter(
        models.Collaborator.project_id == project_id
    ).all()
    return content, citations, owner, collaborators


@router.post("/{project_id}/pdf")
def export_pdf(
    project_id: int,
    style: str = Query(default="apa", enum=["apa", "ieee"]),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = _assert_access(project_id, current_user, db)
    content, citations, owner, collaborators = _load_data(project_id, db)

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=letter,
            leftMargin=1.2 * inch,
            rightMargin=1.2 * inch,
            topMargin=1.2 * inch,
            bottomMargin=1.2 * inch,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title2",
            parent=styles["Normal"],
            fontSize=16,
            fontName="Helvetica-Bold",
            alignment=TA_CENTER,
            spaceAfter=6,
            textColor=colors.HexColor("#1a1a2e"),
        )
        author_style = ParagraphStyle(
            "Authors",
            parent=styles["Normal"],
            fontSize=10,
            fontName="Helvetica",
            alignment=TA_CENTER,
            spaceAfter=4,
            textColor=colors.HexColor("#555555"),
        )
        date_style = ParagraphStyle(
            "Date",
            parent=styles["Normal"],
            fontSize=9,
            fontName="Helvetica-Oblique",
            alignment=TA_CENTER,
            spaceAfter=16,
            textColor=colors.HexColor("#888888"),
        )
        heading_style = ParagraphStyle(
            "Heading2",
            parent=styles["Normal"],
            fontSize=12,
            fontName="Helvetica-Bold",
            spaceBefore=14,
            spaceAfter=4,
            textColor=colors.HexColor("#1a1a2e"),
        )
        body_style = ParagraphStyle(
            "Body2",
            parent=styles["Normal"],
            fontSize=10,
            fontName="Helvetica",
            leading=15,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            textColor=colors.HexColor("#222222"),
        )
        bib_style = ParagraphStyle(
            "Bib",
            parent=styles["Normal"],
            fontSize=9,
            fontName="Helvetica",
            leading=13,
            leftIndent=24,
            firstLineIndent=-24,
            spaceAfter=6,
            textColor=colors.HexColor("#333333"),
        )

        story = []

        story.append(Paragraph(project.title, title_style))

        author_names = [owner.email.split("@")[0]] if owner else []
        for c in collaborators:
            name = c.email.split("@")[0]
            if name not in author_names:
                author_names.append(name)
        story.append(Paragraph(", ".join(author_names), author_style))
        story.append(Paragraph(datetime.utcnow().strftime("%B %d, %Y"), date_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
        story.append(Spacer(1, 12))

        for sec_key in SECTIONS_ORDER:
            text = content.get(sec_key, "").strip()
            if not text:
                continue
            story.append(Paragraph(SECTION_LABELS[sec_key], heading_style))
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe, body_style))

        if citations:
            story.append(Spacer(1, 12))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
            story.append(Paragraph("References", heading_style))
            for i, cit in enumerate(citations, 1):
                if style == "ieee":
                    ref = cit.formatted_ieee or f"[{i}] {cit.title}. {cit.journal or ''}, {cit.year or ''}."
                else:
                    try:
                        raw = cit.authors or "[]"
                        alist = json.loads(raw) if isinstance(raw, str) else (raw if isinstance(raw, list) else [])
                        author_str = ", ".join(str(a) for a in alist) if alist else ""
                    except Exception:
                        author_str = ""
                    ref = cit.formatted_apa or f"{author_str} ({cit.year or 'n.d.'}). {cit.title}. {cit.journal or ''}."
                safe_ref = ref.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_ref, bib_style))

        doc.build(story)
        buf.seek(0)
        safe_title = "".join(c for c in project.title if c.isalnum() or c in " _-").strip()[:50] or "manuscript"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}_{style}.pdf"'},
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}") from exc


@router.post("/{project_id}/docx")
def export_docx(
    project_id: int,
    style: str = Query(default="apa", enum=["apa", "ieee"]),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = _assert_access(project_id, current_user, db)
    content, citations, owner, collaborators = _load_data(project_id, db)

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section_obj in doc.sections:
            section_obj.top_margin = Inches(1.2)
            section_obj.bottom_margin = Inches(1.2)
            section_obj.left_margin = Inches(1.2)
            section_obj.right_margin = Inches(1.2)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(project.title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        author_names = [owner.email.split("@")[0]] if owner else []
        for c in collaborators:
            name = c.email.split("@")[0]
            if name not in author_names:
                author_names.append(name)
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(", ".join(author_names))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(datetime.utcnow().strftime("%B %d, %Y"))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        for sec_key in SECTIONS_ORDER:
            text = content.get(sec_key, "").strip()
            if not text:
                continue
            heading = doc.add_paragraph()
            run = heading.add_run(SECTION_LABELS[sec_key])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            for para_text in text.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = para.add_run(para_text)
                    run.font.size = Pt(10)

        if citations:
            doc.add_paragraph()
            ref_heading = doc.add_paragraph()
            run = ref_heading.add_run("References")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            for i, cit in enumerate(citations, 1):
                if style == "ieee":
                    ref = cit.formatted_ieee or f"[{i}] {cit.title}. {cit.journal or ''}, {cit.year or ''}."
                else:
                    try:
                        raw = cit.authors or "[]"
                        authors_list = json.loads(raw) if isinstance(raw, str) else (raw if isinstance(raw, list) else [])
                        author_str = ", ".join(str(a) for a in authors_list) if authors_list else ""
                    except Exception:
                        author_str = ""
                    ref = cit.formatted_apa or f"{author_str} ({cit.year or 'n.d.'}). {cit.title}. {cit.journal or ''}."
                ref_para = doc.add_paragraph()
                ref_para.paragraph_format.left_indent = Inches(0.3)
                run = ref_para.add_run(ref)
                run.font.size = Pt(9)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_title = "".join(c for c in project.title if c.isalnum() or c in " _-").strip()[:50] or "manuscript"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}_{style}.docx"'},
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}") from exc
